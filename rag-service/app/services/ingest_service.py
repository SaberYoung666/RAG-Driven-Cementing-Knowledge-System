"""实现文档摄取流程，包括 OCR 预处理、清洗、切分与持久化。"""

from __future__ import annotations

from dataclasses import dataclass
import gc
import json
import logging
from pathlib import Path
import re
import tempfile
import time
from typing import Any, Callable, Dict, List, Literal, Tuple
from uuid import uuid4

from app.core.config import PROJECT_ROOT, Settings
from app.core.errors import ServiceError
from app.services.ocr_engine import OCRPageResult, OCREngineUnavailable, PaddleOCREngine
from app.services.ocr_process import OCRProcessConfig, OCRProcessRunner

logger = logging.getLogger(__name__)

EnableOCR = Literal["auto", "on", "off"]
IMG_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
StageCallback = Callable[[str, Dict[str, Any] | None], None]


@dataclass
class PageText:
    page_no: int
    text: str
    ocr_used: bool = False
    ocr_confidence_avg: float | None = None


@dataclass
class DocumentChunks:
    doc_id: str
    source_file_name: str
    chunk_count: int
    pages_processed: int
    ocr_pages: int
    failed_pages: List[int]
    message: str
    elapsed_ms: int
    chunk_file_path: str | None = None
    ocr_text_path: str | None = None


class IngestService:
    def __init__(self, settings: Settings):
        self.settings = settings
        ocr_cfg = settings.config.get("ocr", {})
        self._ocr_engine = PaddleOCREngine(
            lang=str(ocr_cfg.get("lang", "ch")),
            use_angle_cls=bool(ocr_cfg.get("use_angle_cls", True)),
            enable_mkldnn=bool(ocr_cfg.get("enable_mkldnn", False)),
            cpu_threads=(
                int(ocr_cfg["cpu_threads"])
                if ocr_cfg.get("cpu_threads") is not None
                else None
            ),
            init_kwargs=dict(ocr_cfg.get("init_kwargs", {}) or {}),
        )
        self._ocr_runner = OCRProcessRunner(
            OCRProcessConfig(
                lang=self._ocr_engine.lang,
                use_angle_cls=self._ocr_engine.use_angle_cls,
                enable_mkldnn=self._ocr_engine.enable_mkldnn,
                cpu_threads=self._ocr_engine.cpu_threads,
                init_kwargs=dict(self._ocr_engine.init_kwargs or {}),
                prewarm=bool(ocr_cfg.get("prewarm", True)),
            )
        )
        if bool(ocr_cfg.get("prewarm", True)):
            self._ocr_runner.start()

    def close(self) -> None:
        self._ocr_runner.close()

    def run(
            self,
            input_dir: str | None = None,
            strategy: str | None = None,
            incremental: bool = False,
            file_paths: list[str] | None = None,
            enable_ocr: EnableOCR = "auto",
            strict_ocr: bool | None = None,
    ) -> Dict[str, Any]:
        cfg = self.settings.config.get("ingestion", {})
        files = self._resolve_files(
            input_dir=input_dir,
            file_paths=file_paths,
            cfg_input_dir=str(cfg.get("input_dir", "data/raw_docs")),
        )
        if not files:
            if not incremental:
                self.settings.chunks_path.parent.mkdir(parents=True, exist_ok=True)
                self.settings.chunks_path.write_text("", encoding="utf-8")
            return {
                "chunks_path": str(self.settings.chunks_path),
                "files": 0,
                "chunks": 0,
                "total_chunks": 0,
                "chunks_by_source": {},
                "processed_sources": [],
                "processed_doc_ids": [],
                "doc_stats": {},
            }

        docs: List[DocumentChunks] = []
        for path in files:
            doc = self.process_document(
                doc_id=path.name,
                file_path=path,
                source_file_name=path.name,
                strategy=strategy,
                enable_ocr=enable_ocr,
                strict_ocr=strict_ocr,
            )
            docs.append(doc)

        return self.persist_documents(docs=docs, incremental=incremental)

    # 处理文档接口
    def process_document(
            self,
            doc_id: str,
            file_path: Path,
            source_file_name: str,
            strategy: str | None = None,
            enable_ocr: EnableOCR = "auto",
            strict_ocr: bool | None = None,
            status_callback: StageCallback | None = None,
    ) -> DocumentChunks:
        t0 = time.perf_counter()
        cfg = self.settings.config.get("ingestion", {})
        use_strategy = strategy or str(cfg.get("strategy", "structure_then_window"))
        chunk_chars = int(cfg.get("chunk_chars", 1200))
        chunk_overlap = int(cfg.get("chunk_overlap", 200))
        min_chunk_chars = int(cfg.get("min_chunk_chars", 200))
        max_section_chars = int(cfg.get("max_section_chars", 3000))
        remove_empty_lines = bool(cfg.get("remove_empty_lines", True))
        remove_repeated_margins = bool(cfg.get("remove_repeated_margins", True))
        gc_every_pages = 8

        ocr_cfg = self.settings.config.get("ocr", {})
        threshold = int(ocr_cfg.get("min_text_chars", 20))
        strict = (
            bool(ocr_cfg.get("strict", False)) if strict_ocr is None else strict_ocr
        )
        timeout_s = float(ocr_cfg.get("timeout_seconds", 20))

        logger.info(
            "开始处理文档 doc_id=%s file=%s enable_ocr=%s",
            doc_id,
            file_path,
            enable_ocr,
        )
        total_pages = self._estimate_total_pages(file_path)
        extracted_pages_path: Path | None = None
        cleaned_pages_path: Path | None = None
        chunk_file_path: Path | None = None
        ocr_text_path: str | None = None
        pages_processed = 0
        ocr_pages = 0
        failed_pages: List[int] = []
        chunk_count = 0

        try:
            extracted_pages_path, pages_processed, ocr_pages, failed_pages, ocr_text_path = (
                self._extract_pages_to_temp(
                    doc_id=doc_id,
                    file_path=file_path,
                    source_file_name=source_file_name,
                    enable_ocr=enable_ocr,
                    threshold=threshold,
                    strict=strict,
                    timeout_s=timeout_s,
                    status_callback=status_callback,
                )
            )

            if failed_pages:
                logger.warning(
                    "处理文档部分页面失败 doc_id=%s file=%s failed_pages=%s",
                    doc_id,
                    file_path,
                    failed_pages,
                )

            if status_callback is not None:
                status_callback(
                    "cleaning",
                    {
                        "pages_processed": pages_processed,
                        "total_pages": total_pages,
                        "current_page": 0 if pages_processed else None,
                        "stage_progress": 0,
                        "ocr_pages": ocr_pages,
                        "failed_pages": failed_pages,
                    },
                )

            cleaned_pages_path = self._make_temp_path("cleaned-pages-", ".jsonl")
            if remove_repeated_margins and pages_processed > 1:
                logger.info(
                    "低内存路径已禁用整本文档重复页眉页脚检测 doc_id=%s file=%s",
                    doc_id,
                    file_path,
                )
            with cleaned_pages_path.open("w", encoding="utf-8") as cleaned_fh:
                for idx, page_row in enumerate(
                    self._iterate_page_records(extracted_pages_path), 1
                ):
                    page_row.text = self._basic_clean(
                        page_row.text, remove_empty_lines=remove_empty_lines
                    )
                    self._write_page_record(cleaned_fh, page_row)
                    if status_callback is not None and pages_processed > 0:
                        status_callback(
                            "cleaning",
                            {
                                "pages_processed": pages_processed,
                                "total_pages": total_pages,
                                "current_page": idx,
                                "stage_progress": round(idx * 100 / pages_processed),
                                "ocr_pages": ocr_pages,
                                "failed_pages": failed_pages,
                            },
                        )
                    if idx % gc_every_pages == 0:
                        gc.collect()

            if status_callback is not None:
                status_callback(
                    "splitting",
                    {
                        "pages_processed": 0,
                        "total_pages": total_pages,
                        "current_page": 0 if pages_processed else None,
                        "stage_progress": 0,
                        "ocr_pages": ocr_pages,
                        "failed_pages": failed_pages,
                    },
                )

            chunk_file_path = self._make_temp_path("chunks-", ".jsonl")
            with chunk_file_path.open("w", encoding="utf-8") as chunk_fh:
                for idx, page_row in enumerate(
                    self._iterate_page_records(cleaned_pages_path), 1
                ):
                    chunks = self._chunk_page(
                        text=page_row.text,
                        strategy=use_strategy,
                        chunk_chars=chunk_chars,
                        chunk_overlap=chunk_overlap,
                        min_chunk_chars=min_chunk_chars,
                        max_section_chars=max_section_chars,
                    )
                    for chunk_idx, (section, piece) in enumerate(chunks):
                        chunk_fh.write(
                            json.dumps(
                                {
                                    "chunk_id": f"{doc_id}::p{page_row.page_no}::c{chunk_idx}",
                                    "text": piece,
                                    "metadata": {
                                        "doc_id": doc_id,
                                        "source": source_file_name,
                                        "source_file_name": source_file_name,
                                        "page": page_row.page_no,
                                        "page_no": page_row.page_no,
                                        "section": section,
                                        "chunk_index": chunk_idx,
                                        "ocr_used": page_row.ocr_used,
                                        "ocr_confidence_avg": page_row.ocr_confidence_avg,
                                    },
                                },
                                ensure_ascii=False,
                            )
                            + "\n"
                        )
                        chunk_count += 1
                    if status_callback is not None and pages_processed > 0:
                        status_callback(
                            "splitting",
                            {
                                "pages_processed": idx,
                                "total_pages": total_pages,
                                "current_page": idx,
                                "stage_progress": round(idx * 100 / pages_processed),
                                "ocr_pages": ocr_pages,
                                "chunk_count": chunk_count,
                                "failed_pages": failed_pages,
                            },
                        )
                    if idx % gc_every_pages == 0:
                        gc.collect()
        except Exception:
            self._safe_unlink(chunk_file_path)
            self._safe_unlink(ocr_text_path)
            raise
        finally:
            self._safe_unlink(extracted_pages_path)
            self._safe_unlink(cleaned_pages_path)

        elapsed_ms = int((time.perf_counter() - t0) * 1000)
        msg = "已完成"
        if failed_pages:
            msg = f"已完成，部分页面OCR失败({failed_pages})"
        logger.info(
            "处理文档成功 doc_id=%s pages=%s ocr_pages=%s chunks=%s elapsed_ms=%s",
            doc_id,
            pages_processed,
            ocr_pages,
            chunk_count,
            elapsed_ms,
        )
        return DocumentChunks(
            doc_id=doc_id,
            source_file_name=source_file_name,
            chunk_count=chunk_count,
            pages_processed=pages_processed,
            ocr_pages=ocr_pages,
            failed_pages=failed_pages,
            message=msg,
            elapsed_ms=elapsed_ms,
            chunk_file_path=str(chunk_file_path) if chunk_file_path else None,
            ocr_text_path=ocr_text_path,
        )

    def persist_documents(
            self, docs: List[DocumentChunks], incremental: bool = True
    ) -> Dict[str, Any]:
        chunks_path = self.settings.chunks_path
        chunks_path.parent.mkdir(parents=True, exist_ok=True)

        replaced_doc_ids: set[str] = set()
        replaced_sources: set[str] = set()
        chunks_by_source: Dict[str, int] = {}
        doc_stats: Dict[str, Dict[str, Any]] = {}

        for doc in docs:
            replaced_doc_ids.add(doc.doc_id)
            replaced_sources.add(doc.source_file_name)
            chunks_by_source[doc.source_file_name] = chunks_by_source.get(
                doc.source_file_name, 0
            ) + doc.chunk_count
            doc_stats[doc.doc_id] = {
                "pages_processed": doc.pages_processed,
                "ocr_pages": doc.ocr_pages,
                "chunk_count": doc.chunk_count,
                "failed_pages": doc.failed_pages,
                "message": doc.message,
                "elapsed_ms": doc.elapsed_ms,
                "ocr_text_path": doc.ocr_text_path,
            }

        tmp_merged_path = self._make_temp_path("chunks-merged-", ".jsonl")
        total_chunks = 0
        try:
            with tmp_merged_path.open("w", encoding="utf-8") as out:
                if incremental and chunks_path.exists():
                    with chunks_path.open("r", encoding="utf-8") as existing_fh:
                        for line in existing_fh:
                            if not line.strip():
                                continue
                            rec = json.loads(line)
                            metadata = rec.get("metadata") or {}
                            if (
                                str(metadata.get("doc_id") or "") in replaced_doc_ids
                                or str(metadata.get("source") or "") in replaced_sources
                            ):
                                continue
                            out.write(line)
                            total_chunks += 1

                for doc in docs:
                    if not doc.chunk_file_path:
                        continue
                    with Path(doc.chunk_file_path).open("r", encoding="utf-8") as doc_fh:
                        for line in doc_fh:
                            if not line.strip():
                                continue
                            out.write(line)
                            total_chunks += 1

            tmp_merged_path.replace(chunks_path)
        finally:
            self._safe_unlink(tmp_merged_path)
            for doc in docs:
                self._safe_unlink(doc.chunk_file_path)

        return {
            "chunks_path": str(chunks_path),
            "files": len(docs),
            "chunks": sum(doc.chunk_count for doc in docs),
            "total_chunks": total_chunks,
            "chunks_by_source": chunks_by_source,
            "processed_sources": sorted(replaced_sources),
            "processed_doc_ids": sorted(replaced_doc_ids),
            "doc_stats": doc_stats,
        }

    def _extract_pages_to_temp(
            self,
            doc_id: str,
            file_path: Path,
            source_file_name: str,
            enable_ocr: EnableOCR,
            threshold: int,
            strict: bool,
            timeout_s: float,
            status_callback: StageCallback | None = None,
    ) -> Tuple[Path, int, int, List[int], str | None]:
        suffix = file_path.suffix.lower()
        pages_path = self._make_temp_path("extracted-pages-", ".jsonl")
        ocr_path: str | None = None
        pages_processed = 0
        ocr_pages = 0
        failed_pages: List[int] = []
        try:
            with pages_path.open("w", encoding="utf-8") as pages_fh:
                if suffix == ".pdf":
                    pages_processed, ocr_pages, failed_pages, ocr_path = self._extract_pdf_pages_to_temp(
                        doc_id=doc_id,
                        file_path=file_path,
                        source_file_name=source_file_name,
                        enable_ocr=enable_ocr,
                        threshold=threshold,
                        strict=strict,
                        timeout_s=timeout_s,
                        status_callback=status_callback,
                        pages_fh=pages_fh,
                    )
                elif suffix in IMG_SUFFIXES:
                    if enable_ocr == "off":
                        raise ServiceError(
                            "image files require OCR. please set enable_ocr to auto or on.",
                            status_code=400,
                            code="OCR_REQUIRED",
                            details={"file": str(file_path)},
                        )
                    image = self._load_image(file_path)
                    try:
                        if status_callback is not None:
                            status_callback(
                                "ocr_processing",
                                {
                                    "pages_processed": 0,
                                    "total_pages": 1,
                                    "current_page": 1,
                                    "stage_progress": 0,
                                    "ocr_pages": 0,
                                    "failed_pages": [],
                                },
                            )
                        logger.info("文档触发OCR file=%s mode=%s", file_path, enable_ocr)
                        ocr_res = self._ocr_with_timeout(image, timeout_s=timeout_s)
                        page_row = PageText(
                            page_no=1,
                            text=ocr_res.text,
                            ocr_used=True,
                            ocr_confidence_avg=ocr_res.confidence_avg,
                        )
                        ocr_pages = 1
                        if status_callback is not None:
                            status_callback(
                                "ocr_processing",
                                {
                                    "pages_processed": 1,
                                    "total_pages": 1,
                                    "current_page": 1,
                                    "stage_progress": 100,
                                    "ocr_pages": 1,
                                    "failed_pages": [],
                                },
                            )
                    except Exception as exc:
                        if strict:
                            raise ServiceError(
                                "OCR failed on image",
                                status_code=500,
                                code="OCR_IMAGE_FAILED",
                                details={"error": str(exc)},
                            ) from exc
                        failed_pages = [1]
                        if status_callback is not None:
                            status_callback(
                                "ocr_processing",
                                {
                                    "pages_processed": 1,
                                    "total_pages": 1,
                                    "current_page": 1,
                                    "stage_progress": 100,
                                    "ocr_pages": 0,
                                    "failed_pages": [1],
                                },
                            )
                        page_row = PageText(
                            page_no=1,
                            text="",
                            ocr_used=False,
                            ocr_confidence_avg=None,
                        )
                    self._write_page_record(pages_fh, page_row)
                    pages_processed = 1
                    if page_row.ocr_used:
                        ocr_path = self._prepare_ocr_text_path(
                            doc_id=doc_id, source_file_name=source_file_name
                        )
                        with Path(ocr_path).open("w", encoding="utf-8") as ocr_single_fh:
                            self._append_ocr_text_page(
                                ocr_single_fh,
                                page_no=1,
                                text=page_row.text,
                                first_page=True,
                            )
                    del image
                    gc.collect()
                elif suffix in [".txt", ".md"]:
                    text = file_path.read_text(encoding="utf-8", errors="ignore")
                    self._write_page_record(
                        pages_fh, PageText(page_no=1, text=text, ocr_used=False)
                    )
                    pages_processed = 1
                elif suffix in [".doc", ".docx"]:
                    from docx import Document  # type: ignore

                    doc = Document(str(file_path))
                    body = "\n".join([p.text for p in doc.paragraphs if p.text])
                    self._write_page_record(
                        pages_fh, PageText(page_no=1, text=body, ocr_used=False)
                    )
                    pages_processed = 1
                else:
                    raise ServiceError(
                        f"不支持的文件类型: {file_path}",
                        status_code=400,
                        code="UNSUPPORTED_FILE",
                    )
        except Exception:
            self._safe_unlink(pages_path)
            self._safe_unlink(ocr_path)
            raise

        return pages_path, pages_processed, ocr_pages, failed_pages, ocr_path

    def _extract_pdf_pages_to_temp(
            self,
            doc_id: str,
            file_path: Path,
            source_file_name: str,
            enable_ocr: EnableOCR,
            threshold: int,
            strict: bool,
            timeout_s: float,
            status_callback: StageCallback | None,
            pages_fh: Any,
    ) -> Tuple[int, int, List[int], str | None]:
        try:
            import fitz  # type: ignore
        except Exception as exc:
            raise ServiceError(
                "PyMuPDF不可用. 请安装pymupdf.",
                status_code=500,
                code="DEPENDENCY_MISSING",
                details={"dependency": "pymupdf"},
            ) from exc

        failed_pages: List[int] = []
        short_pages: List[int] = []
        ocr_logged = False
        scale = float(self.settings.config.get("ocr", {}).get("pdf_render_scale", 2.0))
        ocr_mode: EnableOCR = (
            enable_ocr if enable_ocr in ("auto", "on", "off") else "auto"
        )
        pages_processed = 0
        ocr_pages = 0
        ocr_path: str | None = None
        ocr_written_pages = 0
        ocr_fh = None

        try:
            with fitz.open(str(file_path)) as pdf:
                total_pages = pdf.page_count
                for idx in range(pdf.page_count):
                    page_no = idx + 1
                    page = pdf.load_page(idx)
                    extracted = (page.get_text("text") or "").strip()
                    need_ocr = ocr_mode == "on" or (
                            ocr_mode == "auto" and len(extracted) < threshold
                    )

                    if ocr_mode == "off" and len(extracted) < threshold:
                        short_pages.append(page_no)

                    if need_ocr:
                        try:
                            render_t0 = time.perf_counter()
                            image = self._render_pdf_page(page, scale=scale)
                            render_ms = (time.perf_counter() - render_t0) * 1000
                            if status_callback is not None:
                                status_callback(
                                    "ocr_processing",
                                    {
                                        "pages_processed": page_no - 1,
                                        "total_pages": total_pages,
                                        "current_page": page_no,
                                        "stage_progress": round((page_no - 1) * 100 / total_pages),
                                        "ocr_pages": ocr_pages,
                                        "failed_pages": failed_pages,
                                    },
                                )
                            if not ocr_logged:
                                logger.info(
                                    "文档触发OCR file=%s mode=%s",
                                    file_path,
                                    ocr_mode,
                                )
                                ocr_logged = True
                            ocr_res = self._ocr_with_timeout(image, timeout_s=timeout_s)
                            text = ocr_res.text.strip() or extracted
                            page_row = PageText(
                                page_no=page_no,
                                text=text,
                                ocr_used=True,
                                ocr_confidence_avg=ocr_res.confidence_avg,
                            )
                            ocr_pages += 1
                            if status_callback is not None:
                                status_callback(
                                    "ocr_processing",
                                    {
                                        "pages_processed": page_no,
                                        "total_pages": total_pages,
                                        "current_page": page_no,
                                        "stage_progress": round(page_no * 100 / total_pages),
                                        "ocr_pages": ocr_pages,
                                        "failed_pages": failed_pages,
                                    },
                                )
                            logger.info(
                                "页面OCR成功 file=%s page=%s render_ms=%.2f ocr_confidence_avg=%.2f extracted_chars=%s ocr_chars=%s",
                                file_path,
                                page_no,
                                render_ms,
                                ocr_res.confidence_avg or 0.0,
                                len(extracted),
                                len(text),
                                extra={
                                    "page_no": page_no,
                                    "render_ms": round(render_ms, 2),
                                    "ocr_timing": ocr_res.timing or {},
                                },
                            )
                        except Exception as exc:
                            failed_pages.append(page_no)
                            logger.exception(
                                "第%s页OCR失败，原因：%s",
                                page_no,
                                str(exc),
                                extra={"file": str(file_path), "page": page_no},
                            )
                            if strict:
                                raise ServiceError(
                                    f"第{page_no}页OCR失败",
                                    status_code=500,
                                    code="OCR_PAGE_FAILED",
                                    details={"page_no": page_no, "error": str(exc)},
                                ) from exc
                            page_row = PageText(
                                page_no=page_no,
                                text=extracted,
                                ocr_used=False,
                                ocr_confidence_avg=None,
                            )
                            if status_callback is not None:
                                status_callback(
                                    "ocr_processing",
                                    {
                                        "pages_processed": page_no,
                                        "total_pages": total_pages,
                                        "current_page": page_no,
                                        "stage_progress": round(page_no * 100 / total_pages),
                                        "ocr_pages": ocr_pages,
                                        "failed_pages": failed_pages,
                                    },
                                )
                        finally:
                            del page
                            if "image" in locals():
                                del image
                            if "ocr_res" in locals():
                                del ocr_res
                            if page_no % 8 == 0:
                                gc.collect()
                    else:
                        page_row = PageText(
                            page_no=page_no,
                            text=extracted,
                            ocr_used=False,
                            ocr_confidence_avg=None,
                        )
                        if status_callback is not None:
                            status_callback(
                                "processing",
                                {
                                    "pages_processed": page_no,
                                    "total_pages": total_pages,
                                    "current_page": page_no,
                                    "stage_progress": round(page_no * 100 / total_pages),
                                    "ocr_pages": ocr_pages,
                                    "failed_pages": failed_pages,
                                },
                            )
                        del page

                    self._write_page_record(pages_fh, page_row)
                    pages_processed += 1
                    if page_row.ocr_used:
                        if ocr_fh is None:
                            ocr_path = self._prepare_ocr_text_path(
                                doc_id=doc_id, source_file_name=source_file_name
                            )
                            ocr_fh = Path(ocr_path).open("w", encoding="utf-8")
                        self._append_ocr_text_page(
                            ocr_fh,
                            page_no=page_row.page_no,
                            text=page_row.text,
                            first_page=ocr_written_pages == 0,
                        )
                        ocr_written_pages += 1

            if ocr_mode == "off" and short_pages:
                raise ServiceError(
                    "一些页面文本较少，可能需要OCR才能正确提取。请设置enable_ocr为auto或on重试。",
                    status_code=400,
                    code="OCR_REQUIRED",
                    details={"pages": short_pages},
                )
        finally:
            if ocr_fh is not None:
                ocr_fh.close()

        return pages_processed, ocr_pages, failed_pages, ocr_path

    @staticmethod
    def _write_page_record(fh: Any, page: PageText) -> None:
        fh.write(
            json.dumps(
                {
                    "page_no": page.page_no,
                    "text": page.text,
                    "ocr_used": page.ocr_used,
                    "ocr_confidence_avg": page.ocr_confidence_avg,
                },
                ensure_ascii=False,
            )
            + "\n"
        )

    @staticmethod
    def _iterate_page_records(path: Path):
        with path.open("r", encoding="utf-8") as fh:
            for line in fh:
                if not line.strip():
                    continue
                payload = json.loads(line)
                yield PageText(
                    page_no=int(payload.get("page_no", 0) or 0),
                    text=str(payload.get("text", "") or ""),
                    ocr_used=bool(payload.get("ocr_used", False)),
                    ocr_confidence_avg=payload.get("ocr_confidence_avg"),
                )

    def _prepare_ocr_text_path(self, doc_id: str, source_file_name: str) -> str:
        ocr_cfg = self.settings.config.get("ocr", {})
        output_dir = PROJECT_ROOT / str(ocr_cfg.get("output_dir", "data/ocr_texts"))
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_name = self._sanitize_filename(doc_id) or self._sanitize_filename(
            source_file_name
        )
        return str(output_dir / f"{safe_name or uuid4().hex}.txt")

    @staticmethod
    def _append_ocr_text_page(
            fh: Any, page_no: int, text: str, first_page: bool = False
    ) -> None:
        if not first_page:
            fh.write("\n\n")
        header = f"[Page {page_no}]"
        body = (text or "").strip()
        fh.write(f"{header}\n{body}" if body else header)

    def _make_temp_path(self, prefix: str, suffix: str) -> Path:
        tmp_dir = self.settings.chunks_path.parent / "tmp"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        with tempfile.NamedTemporaryFile(
                delete=False,
                dir=str(tmp_dir),
                prefix=prefix,
                suffix=suffix,
        ) as tmp:
            return Path(tmp.name)

    @staticmethod
    def _safe_unlink(path: str | Path | None) -> None:
        if not path:
            return
        try:
            Path(path).unlink(missing_ok=True)
        except Exception:
            pass

    # 基本清洗，包括ocr预处理、去除空行等
    def _extract_pages(
            self,
            file_path: Path,
            enable_ocr: EnableOCR,
            threshold: int,
            strict: bool,
            timeout_s: float,
            status_callback: StageCallback | None = None,
    ) -> Tuple[List[PageText], List[int]]:
        # 获取文件后缀
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf_pages(
                file_path=file_path,
                enable_ocr=enable_ocr,
                threshold=threshold,
                strict=strict,
                timeout_s=timeout_s,
                status_callback=status_callback,
            )
        if suffix in IMG_SUFFIXES:
            return self._extract_image_page(
                file_path=file_path,
                enable_ocr=enable_ocr,
                strict=strict,
                timeout_s=timeout_s,
                status_callback=status_callback,
            )
        if suffix in [".txt", ".md"]:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return [PageText(page_no=1, text=text, ocr_used=False)], []
        if suffix in [".doc", ".docx"]:
            from docx import Document  # type: ignore

            doc = Document(str(file_path))
            body = "\n".join([p.text for p in doc.paragraphs if p.text])
            return [PageText(page_no=1, text=body, ocr_used=False)], []
        raise ServiceError(
            f"不支持的文件类型: {file_path}", status_code=400, code="UNSUPPORTED_FILE"
        )

    # pdf文档提取
    def _extract_pdf_pages(
            self,
            file_path: Path,
            enable_ocr: EnableOCR,
            threshold: int,
            strict: bool,
            timeout_s: float,
            status_callback: StageCallback | None = None,
    ) -> Tuple[List[PageText], List[int]]:
        try:
            import fitz
        except Exception as exc:
            raise ServiceError(
                "PyMuPDF不可用. 请安装pymupdf.",
                status_code=500,
                code="DEPENDENCY_MISSING",
                details={"dependency": "pymupdf"},
            ) from exc

        rows: List[PageText] = []
        failed_pages: List[int] = []
        short_pages: List[int] = []
        ocr_logged = False
        scale = float(self.settings.config.get("ocr", {}).get("pdf_render_scale", 2.0))
        ocr_mode: EnableOCR = (
            enable_ocr if enable_ocr in ("auto", "on", "off") else "auto"
        )

        with fitz.open(str(file_path)) as pdf:
            total_pages = pdf.page_count
            ocr_count = 0
            for idx in range(pdf.page_count):
                page_no = idx + 1
                page = pdf.load_page(idx)
                extracted = (page.get_text("text") or "").strip()
                need_ocr = ocr_mode == "on" or (
                        ocr_mode == "auto" and len(extracted) < threshold
                )

                if ocr_mode == "off" and len(extracted) < threshold:
                    short_pages.append(page_no)

                if need_ocr:
                    try:
                        # 将PDF页转化为图像
                        render_t0 = time.perf_counter()
                        image = self._render_pdf_page(page, scale=scale)
                        render_ms = (time.perf_counter() - render_t0) * 1000
                        if status_callback is not None:
                            status_callback(
                                "ocr_processing",
                                {
                                    "pages_processed": page_no - 1,
                                    "total_pages": total_pages,
                                    "current_page": page_no,
                                    "stage_progress": round((page_no - 1) * 100 / total_pages),
                                    "ocr_pages": ocr_count,
                                    "failed_pages": failed_pages,
                                },
                            )
                        if not ocr_logged:
                            logger.info(
                                "文档触发OCR file=%s mode=%s",
                                file_path,
                                ocr_mode,
                            )
                            ocr_logged = True
                        # 执行OCR逻辑
                        ocr_res = self._ocr_with_timeout(image, timeout_s=timeout_s)
                        text = ocr_res.text.strip() or extracted
                        rows.append(
                            PageText(
                                page_no=page_no,
                                text=text,
                                ocr_used=True,
                                ocr_confidence_avg=ocr_res.confidence_avg,
                            )
                        )
                        ocr_count += 1
                        if status_callback is not None:
                            status_callback(
                                "ocr_processing",
                                {
                                    "pages_processed": page_no,
                                    "total_pages": total_pages,
                                    "current_page": page_no,
                                    "stage_progress": round(page_no * 100 / total_pages),
                                    "ocr_pages": ocr_count,
                                    "failed_pages": failed_pages,
                                },
                            )
                        # TODO：当前调试时，会打印整个页面的内容，之后可以改为只打印OCR提取的文本长度和置信度等信息
                        logger.info(
                            "页面OCR成功 file=%s page=%s render_ms=%.2f ocr_confidence_avg=%.2f extracted_chars=%s ocr_chars=%s",
                            file_path,
                            page_no,
                            render_ms,
                            ocr_res.confidence_avg or 0.0,
                            len(extracted),
                            len(text),
                            extra={
                                "page_no": page_no,
                                "render_ms": round(render_ms, 2),
                                "ocr_timing": ocr_res.timing or {},
                            },
                        )
                    except Exception as exc:
                        failed_pages.append(page_no)
                        logger.exception(
                            "第%s页OCR失败，原因：%s",
                            page_no,
                            str(exc),
                            extra={"file": str(file_path), "page": page_no},
                        )
                        if strict:
                            raise ServiceError(
                                f"第{page_no}页OCR失败",
                                status_code=500,
                                code="OCR_PAGE_FAILED",
                                details={"page_no": page_no, "error": str(exc)},
                            ) from exc
                        rows.append(
                            PageText(
                                page_no=page_no,
                                text=extracted,
                                ocr_used=False,
                                ocr_confidence_avg=None,
                            )
                        )
                        if status_callback is not None:
                            status_callback(
                                "ocr_processing",
                                {
                                    "pages_processed": page_no,
                                    "total_pages": total_pages,
                                    "current_page": page_no,
                                    "stage_progress": round(page_no * 100 / total_pages),
                                    "ocr_pages": ocr_count,
                                    "failed_pages": failed_pages,
                                },
                            )
                else:
                    rows.append(
                        PageText(
                            page_no=page_no,
                            text=extracted,
                            ocr_used=False,
                            ocr_confidence_avg=None,
                        )
                    )
                    if status_callback is not None:
                        status_callback(
                            "processing",
                            {
                                "pages_processed": page_no,
                                "total_pages": total_pages,
                                "current_page": page_no,
                                "stage_progress": round(page_no * 100 / total_pages),
                                "ocr_pages": ocr_count,
                                "failed_pages": failed_pages,
                            },
                        )

        if ocr_mode == "off" and short_pages:
            raise ServiceError(
                "一些页面文本较少，可能需要OCR才能正确提取。请设置enable_ocr为auto或on重试。",
                status_code=400,
                code="OCR_REQUIRED",
                details={"pages": short_pages},
            )
        return rows, failed_pages

    def _extract_image_page(
            self,
            file_path: Path,
            enable_ocr: EnableOCR,
            strict: bool,
            timeout_s: float,
            status_callback: StageCallback | None = None,
    ) -> Tuple[List[PageText], List[int]]:
        if enable_ocr == "off":
            raise ServiceError(
                "image files require OCR. please set enable_ocr to auto or on.",
                status_code=400,
                code="OCR_REQUIRED",
                details={"file": str(file_path)},
            )
        image = self._load_image(file_path)
        try:
            if status_callback is not None:
                status_callback(
                    "ocr_processing",
                    {
                        "pages_processed": 0,
                        "total_pages": 1,
                        "current_page": 1,
                        "stage_progress": 0,
                        "ocr_pages": 0,
                        "failed_pages": [],
                    },
                )
            logger.info("文档触发OCR file=%s mode=%s", file_path, enable_ocr)
            ocr_res = self._ocr_with_timeout(image, timeout_s=timeout_s)
            if status_callback is not None:
                status_callback(
                    "ocr_processing",
                    {
                        "pages_processed": 1,
                        "total_pages": 1,
                        "current_page": 1,
                        "stage_progress": 100,
                        "ocr_pages": 1,
                        "failed_pages": [],
                    },
                )
            return [
                PageText(
                    page_no=1,
                    text=ocr_res.text,
                    ocr_used=True,
                    ocr_confidence_avg=ocr_res.confidence_avg,
                )
            ], []
        except Exception as exc:
            if strict:
                raise ServiceError(
                    "OCR failed on image",
                    status_code=500,
                    code="OCR_IMAGE_FAILED",
                    details={"error": str(exc)},
                ) from exc
            if status_callback is not None:
                status_callback(
                    "ocr_processing",
                    {
                        "pages_processed": 1,
                        "total_pages": 1,
                        "current_page": 1,
                        "stage_progress": 100,
                        "ocr_pages": 0,
                        "failed_pages": [1],
                    },
                )
            return [
                PageText(page_no=1, text="", ocr_used=False, ocr_confidence_avg=None)
            ], [1]

    @staticmethod
    def _estimate_total_pages(file_path: Path) -> int:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            try:
                import fitz
                with fitz.open(str(file_path)) as pdf:
                    return max(1, int(pdf.page_count))
            except Exception:
                return 1
        return 1

    # 实际OCR功能接口
    def _ocr_with_timeout(self, image: Any, timeout_s: float) -> OCRPageResult:
        try:
            t0 = time.perf_counter()
            result = self._ocr_runner.recognize(image=image, timeout_s=timeout_s)
            elapsed_ms = (time.perf_counter() - t0) * 1000
            logger.info(
                "OCR任务完成 backend=process_worker elapsed_ms=%.2f timeout_s=%.2f",
                elapsed_ms,
                timeout_s,
                extra={
                    "ocr_backend": "process_worker",
                    "ocr_worker_exec_ms": round(elapsed_ms, 2),
                    "timeout_s": timeout_s,
                    "ocr_timing": result.timing or {},
                },
            )
            return result
        except TimeoutError:
            raise
        except OCREngineUnavailable as exc:
            logger.error("OCR引擎不可用，无法执行OCR", extra={"error": str(exc)})
            raise ServiceError(
                "PaddleOCR不可用。请安装paddleocr和paddlepaddle。",
                status_code=500,
                code="DEPENDENCY_MISSING",
                details={"dependency": "paddleocr/paddlepaddle", "error": str(exc)},
            ) from exc
        except Exception as exc:
            raise RuntimeError(f"ocr execution failed: {exc}") from exc

    @staticmethod
    def _load_image(path: Path):
        try:
            import numpy as np  # type: ignore
            from PIL import Image  # type: ignore
        except Exception as exc:
            raise ServiceError(
                "Image OCR dependencies are unavailable. Please install pillow and numpy.",
                status_code=500,
                code="DEPENDENCY_MISSING",
                details={"dependency": "pillow/numpy"},
            ) from exc
        img = Image.open(path).convert("RGB")
        return np.array(img)

    # 渲染PDF页面为图像
    @staticmethod
    def _render_pdf_page(page: Any, scale: float):
        try:
            import numpy as np  # type: ignore
            import fitz  # type: ignore
        except Exception as exc:
            raise ServiceError(
                "渲染PDF的依赖不可用。请安装 pymupdf 和 numpy。",
                status_code=500,
                code="DEPENDENCY_MISSING",
                details={"dependency": "pymupdf/numpy"},
            ) from exc
        matrix = fitz.Matrix(scale, scale)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, pix.n
        ).copy()
        del pix
        return arr

    def _chunk_page(
            self,
            text: str,
            strategy: str,
            chunk_chars: int,
            chunk_overlap: int,
            min_chunk_chars: int,
            max_section_chars: int,
    ) -> List[Tuple[str | None, str]]:
        text = text.strip()
        if not text:
            return []
        chunks: List[Tuple[str | None, str]] = []
        if strategy == "structure_then_window":
            for sec_title, sec_text in self._split_into_sections(text):
                if len(sec_text) > max_section_chars:
                    for piece in self._sliding(
                            sec_text, chunk_chars, chunk_overlap, min_chunk_chars
                    ):
                        chunks.append((sec_title, piece))
                elif len(sec_text) >= min_chunk_chars:
                    chunks.append((sec_title, sec_text))
            return chunks
        for piece in self._sliding(text, chunk_chars, chunk_overlap, min_chunk_chars):
            chunks.append((None, piece))
        return chunks

    @staticmethod
    def _remove_repeated_margin_lines(
            rows: List[PageText], top_n: int = 2, bottom_n: int = 2
    ) -> List[PageText]:
        candidates: Dict[str, int] = {}
        for row in rows:
            lines = [x.strip() for x in row.text.split("\n") if x.strip()]
            for ln in lines[:top_n] + lines[-bottom_n:]:
                if 1 < len(ln) <= 80:
                    candidates[ln] = candidates.get(ln, 0) + 1
        threshold = max(2, len(rows) // 2)
        repeated = {ln for ln, cnt in candidates.items() if cnt >= threshold}
        if not repeated:
            return rows
        out: List[PageText] = []
        for row in rows:
            kept = [
                x
                for x in row.text.split("\n")
                if x.strip() and x.strip() not in repeated
            ]
            out.append(
                PageText(
                    page_no=row.page_no,
                    text="\n".join(kept),
                    ocr_used=row.ocr_used,
                    ocr_confidence_avg=row.ocr_confidence_avg,
                )
            )
        return out

    def _persist_ocr_text_document(
            self, doc_id: str, source_file_name: str, pages: List[PageText]
    ) -> str:
        ocr_cfg = self.settings.config.get("ocr", {})
        output_dir = PROJECT_ROOT / str(ocr_cfg.get("output_dir", "data/ocr_texts"))
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_name = self._sanitize_filename(doc_id) or self._sanitize_filename(
            source_file_name
        )
        out_path = output_dir / f"{safe_name or uuid4().hex}.txt"
        out_path.write_text(self._compose_document_text(pages), encoding="utf-8")
        logger.info("已保存OCR文本 doc_id=%s path=%s", doc_id, out_path)
        return str(out_path)

    @staticmethod
    def _compose_document_text(pages: List[PageText]) -> str:
        parts: List[str] = []
        for row in pages:
            header = f"[Page {row.page_no}]"
            body = (row.text or "").strip()
            parts.append(f"{header}\n{body}" if body else header)
        return "\n\n".join(parts).strip()

    @staticmethod
    def _sanitize_filename(name: str) -> str:
        sanitized = re.sub(r'[<>:"/\\|?*]+', "_", name or "").strip(" .")
        return sanitized or "ocr_document"

    def _resolve_files(
            self, input_dir: str | None, file_paths: list[str] | None, cfg_input_dir: str
    ) -> List[Path]:
        if file_paths:
            files = [self._to_abs_path(Path(p)) for p in file_paths]
            return sorted([p for p in files if p.exists() and p.is_file()])
        source_dir = PROJECT_ROOT / str(input_dir or cfg_input_dir)
        return self._collect_files(source_dir)

    @staticmethod
    def _to_abs_path(path: Path) -> Path:
        if path.is_absolute():
            return path
        return (PROJECT_ROOT / path).resolve()

    @staticmethod
    def _read_existing_chunks(path: Path) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        with path.open("r", encoding="utf-8") as fh:
            for line in fh:
                if not line.strip():
                    continue
                out.append(json.loads(line))
        return out

    @staticmethod
    def _collect_files(input_dir: Path) -> List[Path]:
        files: List[Path] = []
        for ext in (
                "*.pdf",
                "*.docx",
                "*.txt",
                "*.md",
                "*.png",
                "*.jpg",
                "*.jpeg",
                "*.webp",
                "*.bmp",
        ):
            files.extend(input_dir.rglob(ext))
        return sorted(files)

    #
    @staticmethod
    def _basic_clean(text: str, remove_empty_lines: bool = True) -> str:
        text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        if not remove_empty_lines:
            return text.strip()
        return "\n".join([x.strip() for x in text.split("\n") if x.strip()])

    @staticmethod
    def _split_into_sections(text: str) -> List[Tuple[str, str]]:
        lines = text.split("\n")
        blocks: List[Tuple[str, str]] = []
        current = "正文"
        acc: List[str] = []
        pat = re.compile(
            r"^\s*(\d+(\.\d+){0,4}\s+.+|[一二三四五六七八九十]+、\s*.+|（[一二三四五六七八九十]+）\s*.+|第\d+[章节].*)$"
        )
        for ln in lines:
            if pat.match(ln.strip()):
                if acc:
                    blocks.append((current, "\n".join(acc).strip()))
                current = ln.strip()
                acc = []
            else:
                acc.append(ln)
        if acc:
            blocks.append((current, "\n".join(acc).strip()))
        return [b for b in blocks if b[1]] or [("正文", text.strip())]

    @staticmethod
    def _sliding(
            text: str, chunk_chars: int, chunk_overlap: int, min_chunk_chars: int
    ) -> List[str]:
        text = text.strip()
        if not text:
            return []
        if chunk_overlap >= chunk_chars:
            raise ValueError("chunk_overlap must be less than chunk_chars")
        if len(text) <= chunk_chars:
            return [text] if len(text) >= min_chunk_chars else []

        out: List[str] = []
        start = 0
        step = chunk_chars - chunk_overlap
        while start < len(text):
            end = min(len(text), start + chunk_chars)
            part = text[start:end].strip()
            if len(part) >= min_chunk_chars:
                out.append(part)
            if end == len(text):
                break
            start += step
        return out

    @staticmethod
    def make_doc_id() -> str:
        return uuid4().hex
